import json
from pathlib import Path
from typing import Any, Dict, List

from app.core.logging_utils import get_logger
from app.services.ingestion.mineru import (
    extract_pdf_text_with_mineru,
    mineru_available,
)
from app.services.ingestion.records import normalize_items
from app.services.ingestion.text_utils import (
    build_parent_id,
    build_source_title,
    build_summary,
    clean_name,
    clean_text,
)


logger = get_logger(__name__)
MAX_INGESTION_TEXT_LENGTH = 60000


def load_json_records(file_path: Path) -> List[Dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8-sig") as file:
        data = json.load(file)
    if isinstance(data, dict):
        candidate = data.get("data") or data.get("items") or data.get("records")
        if isinstance(candidate, list):
            data = candidate
    if not isinstance(data, list):
        raise ValueError(f"{file_path.name} must contain a JSON array")
    return normalize_items(data, file_path.name, source_file_name=file_path.name)


def read_text_with_fallbacks(file_path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"{file_path.name} must use UTF-8 or GB18030 compatible encoding")


def read_pdf_text(file_path: Path) -> str:
    if mineru_available():
        try:
            logger.info("开始使用 MinerU 解析 PDF: 文件=%s", file_path.name)
            text = extract_pdf_text_with_mineru(file_path)
            if text:
                logger.info("MinerU PDF 解析成功: 文件=%s", file_path.name)
                return text
            logger.warning("MinerU PDF 解析结果为空，回退本地解析: 文件=%s", file_path.name)
        except Exception:
            logger.exception("MinerU PDF 解析失败，回退本地解析: 文件=%s", file_path.name)

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency pypdf, please install requirements.txt"
        ) from exc

    logger.info("开始使用本地 pypdf 解析 PDF: 文件=%s", file_path.name)
    reader = PdfReader(str(file_path))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def read_docx_text(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency python-docx, please install requirements.txt"
        ) from exc
    document = Document(str(file_path))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def _finalize_ingestion_text(file_path: Path, raw_text: str) -> str:
    message = clean_text(raw_text)
    if len(message) > MAX_INGESTION_TEXT_LENGTH:
        logger.warning(
            "解析后的文本过长，已截断后再入库: 文件=%s 原始长度=%s 截断长度=%s",
            file_path.name,
            len(message),
            MAX_INGESTION_TEXT_LENGTH,
        )
        message = message[:MAX_INGESTION_TEXT_LENGTH].rstrip()
    return clean_text(message)


def load_plain_text_records(file_path: Path) -> List[Dict[str, Any]]:
    suffix = file_path.suffix.lower()
    if suffix == ".txt":
        raw_text = read_text_with_fallbacks(file_path)
    elif suffix == ".pdf":
        raw_text = read_pdf_text(file_path)
    elif suffix == ".docx":
        raw_text = read_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    message = _finalize_ingestion_text(file_path, raw_text)
    if not message:
        raise ValueError(f"{file_path.name} does not contain usable text")

    file_stem_name = clean_name(file_path.stem) or "文档"
    return [
        {
            "name": file_stem_name,
            "message": message,
            "summary": build_summary(message),
            "parent_id": build_parent_id(file_stem_name, file_path.name, message),
            "source_title": build_source_title(file_path.name),
            "aliases": [],
            "source_file": file_path.name,
        }
    ]
